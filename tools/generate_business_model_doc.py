from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "MODELO_NEGOCIO_NOVA_FORMA_CRM.md"
DOCX_PATH = DOCS / "Modelo_Negocio_Nova_Forma_CRM.docx"


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
GOLD = RGBColor(183, 122, 35)
GRAY = RGBColor(82, 91, 106)
LIGHT = "F4F6F9"
PALE_GOLD = "FFF4DF"
PALE_BLUE = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, bold=False,
                  italic=False, color=None, size=None, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
    return p


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(35, 42, 54)
    add_paragraph(doc, "", after=4)


def add_table(doc: Document, headers, rows, widths=None, header_fill=PALE_BLUE):
    widths = widths or [int(9360 / len(headers))] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = NAVY
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(35, 42, 54)
    add_paragraph(doc, "", after=5)
    return table


SECTIONS = [
    ("1. Sumario executivo", [
        "O Nova Forma CRM deve ser posicionado como um CRM vertical para construtoras, empresas de steel frame, reformas e obras sob medida que recebem demanda pelo WhatsApp, site, Google Meu Negocio, Instagram e indicacoes.",
        "A oportunidade nao esta em competir com CRMs genericos. O valor esta em transformar conversas soltas em processo comercial: lead cadastrado, qualificacao, follow-up, visita tecnica, briefing, orcamento, negociacao e fechamento.",
        "A tese central e vender resultado comercial: menos leads perdidos, mais visitas marcadas, mais orcamentos enviados e maior controle da carteira de oportunidades."
    ]),
    ("2. Produto e proposta de valor", [
        "O aplicativo organiza toda a jornada comercial de uma construtora pequena ou media: captura de leads, funil, tarefas, interacoes, templates de WhatsApp, IA para preenchimento de lead por conversa ou print, dashboard, painel de parceiro tecnico e briefing de visita.",
        "A promessa comercial deve ser simples: 'organize seu WhatsApp comercial e pare de perder oportunidades de obra'.",
        "O produto substitui planilhas, blocos de notas, historicos soltos de WhatsApp e memoria do vendedor por um fluxo rastreavel e acionavel."
    ]),
    ("3. Cliente ideal", [
        "O ICP inicial deve ser construtora de pequeno e medio porte que vende obras residenciais, steel frame, ampliacoes, reformas, casas modulares ou construcao a seco, com ticket medio alto e ciclo comercial consultivo.",
        "Sinais de cliente ideal: recebe leads pelo WhatsApp, tem mais de um atendimento por semana, faz visita tecnica antes do orcamento, perde contato por falta de follow-up, usa planilha e ainda nao possui um CRM adaptado ao processo de obra.",
        "Clientes menos prioritarios no inicio: empresas com baixa recorrencia de leads, ticket muito baixo, operacoes que exigem ERP completo ou companhias que ja possuem CRM corporativo fortemente integrado."
    ]),
    ("4. Problemas que o app resolve", [
        "Lead chega pelo WhatsApp e nao vira cadastro estruturado.",
        "Vendedor esquece retorno depois do primeiro contato.",
        "Cliente pergunta preco por metro quadrado, bairro, visita ou prazo, mas a conversa nao entra no funil.",
        "Dono da empresa nao sabe quantos leads estao quentes, atrasados, em orcamento ou prontos para visita.",
        "Parceiro tecnico faz visita, mas o retorno nao fica centralizado.",
        "A equipe nao sabe quais mensagens usar em cada etapa."
    ]),
    ("5. Diferenciais competitivos", [
        "CRM generico exige configuracao e treinamento. O Nova Forma CRM ja nasce com linguagem, campos, funil e rotinas de construcao.",
        "A IA de importacao por conversa ou print reduz atrito no cadastro. O usuario cola uma conversa, envia prints e revisa um rascunho antes de salvar.",
        "A area de parceiro permite separar o trabalho comercial do retorno tecnico sem perder controle do lead.",
        "O foco em WhatsApp e templates torna o uso diario mais natural para pequenas construtoras do que ferramentas de vendas corporativas."
    ]),
    ("6. Oferta comercial", [
        "A oferta inicial recomendada e vender o app como implantacao assistida mais assinatura mensal.",
        "Mensagem de oferta: 'Em 7 dias organizamos seu atendimento comercial: funil, leads, follow-ups, templates de WhatsApp, dashboard e rotina de visitas'.",
        "Essa embalagem evita que o cliente compare apenas mensalidade de software e permite capturar valor pelo setup, treinamento e configuracao inicial."
    ]),
    ("7. Modelo de receita", [
        "O modelo principal deve ser assinatura SaaS mensal, com implantacao cobrada a parte. A mensalidade financia produto, infraestrutura, suporte e evolucao. O setup paga configuracao, migracao e treinamento.",
        "Tambem podem existir receitas adicionais: importacao de base antiga, configuracao de dominio/app, personalizacao de templates, treinamento comercial e consultoria de processo.",
        "No futuro, envio automatizado via WhatsApp API deve ser cobrado como add-on, pois pode gerar custo variavel por mensagem e exigencias de compliance."
    ]),
    ("8. Estrategia de go-to-market", [
        "Comecar com nicho estreito: empresas de steel frame e construcao a seco. A dor e muito clara, o vocabulario e especifico e o produto ja foi construidodo a partir desse contexto.",
        "Canais iniciais: venda direta para construtoras locais, grupos de steel frame, parcerias com engenheiros, arquitetos e fornecedores, conteudo mostrando organizacao de WhatsApp, anuncios segmentados e indicacoes.",
        "A demonstracao deve usar conversas reais anonimizadas: antes, lead solto no WhatsApp; depois, lead estruturado, follow-up criado, template enviado e visita agendada."
    ]),
    ("9. Operacao e entrega", [
        "A operacao precisa ser leve e repetivel: diagnostico inicial, importacao de leads, configuracao de funil, treinamento do dono e equipe, ativacao de templates, configuracao do parceiro tecnico e revisao apos 15 dias.",
        "O sucesso do cliente deve ser medido por uso real: leads cadastrados, tarefas criadas, interacoes registradas, visitas marcadas e orcamentos enviados.",
        "O suporte inicial pode ser via WhatsApp e video curto. Depois, criar base de conhecimento, checklists e videos de treinamento."
    ]),
    ("10. Metricas de gestao", [
        "Metricas comerciais do cliente: total de leads, novos leads, leads qualificados, visitas marcadas, orcamentos enviados, negociacao, obras fechadas, leads perdidos, leads quentes, leads atrasados, valor potencial e taxa de conversao.",
        "Metricas do negocio SaaS: MRR, churn, CAC, LTV, taxa de ativacao, tempo ate primeiro lead cadastrado, uso semanal ativo, ticket medio, expansao por usuarios/parceiros e receita de setup.",
        "Metricas de produto: tempo para cadastrar lead, taxa de leads com proxima tarefa, uso de templates, uso da IA, falhas na importacao e quantidade de leads sem retorno."
    ]),
    ("11. Riscos e mitigacoes", [
        "Risco de baixa disciplina de uso: mitigar com onboarding pratico, templates prontos e dashboard simples.",
        "Risco de depender demais da IA: manter revisao obrigatoria antes de salvar e orientar conferirmos nome, telefone e cidade.",
        "Risco de concorrencia com CRMs grandes: reforcar nicho, implantacao assistida e fluxo especifico de obra.",
        "Risco de custo operacional alto: padronizar onboarding, criar materiais de treinamento e limitar customizacoes fora dos planos."
    ]),
    ("12. Roadmap estrategico", [
        "Fase 1: estabilizar produto, vender para usuarios proximos e validar preco.",
        "Fase 2: criar onboarding padrao, paginas comerciais, demonstracao guiada e playbook de vendas.",
        "Fase 3: adicionar recursos avancados somente apos tracao: WhatsApp API oficial, relatorios por vendedor, automacoes, multiempresa e marketplace de parceiros.",
        "Fase 4: transformar conhecimento vertical em barreira competitiva: benchmarks de conversao, biblioteca de roteiros, checklist tecnico e inteligencia por tipo de obra."
    ]),
]


TABLES = {
    "pricing": {
        "headers": ["Plano", "Preco sugerido", "Publico", "Inclui"],
        "rows": [
            ["Essencial", "R$ 97/mes", "Autonomos e construtoras pequenas", "1 usuario, funil, tarefas, leads limitados, templates e dashboard basico"],
            ["Profissional", "R$ 197/mes", "Construtoras em operacao comercial ativa", "3 usuarios, leads ilimitados, IA, interacoes, parceiro, briefing e dashboard completo"],
            ["Empresa", "R$ 397/mes", "Equipes maiores e operacoes com parceiros", "10 usuarios, permissoes, parceiros, suporte prioritario e relatorios avancados"],
            ["Setup", "R$ 497 a R$ 1.500", "Todos os planos", "Configuracao, importacao, treinamento, funil e templates iniciais"],
        ],
        "widths": [1800, 1700, 2600, 3260],
    },
    "journey": {
        "headers": ["Etapa", "Dor atual", "Como o CRM resolve", "Resultado esperado"],
        "rows": [
            ["Chegada do lead", "Contato fica perdido no WhatsApp", "IA ou cadastro manual cria lead estruturado", "Lead entra no funil rapidamente"],
            ["Qualificacao", "Faltam cidade, terreno, planta e prazo", "Campos obrigatorios e perguntas sugeridas", "Atendimento mais consultivo"],
            ["Follow-up", "Retorno depende da memoria do vendedor", "Tarefas e alertas de atraso", "Menos oportunidades esquecidas"],
            ["Visita tecnica", "Briefing fica solto", "Resumo do lead e retorno do parceiro", "Orcamento com mais contexto"],
            ["Orcamento", "Cliente some apos proposta", "Templates e follow-ups por etapa", "Melhor taxa de conversao"],
        ],
        "widths": [1700, 2500, 2700, 2460],
    },
    "financial": {
        "headers": ["Cenario", "Clientes pagantes", "Mensalidade media", "MRR estimado", "Setup medio"],
        "rows": [
            ["Piloto", "5", "R$ 197", "R$ 985", "R$ 997"],
            ["Validacao", "20", "R$ 197", "R$ 3.940", "R$ 997"],
            ["Tracao local", "50", "R$ 247", "R$ 12.350", "R$ 1.200"],
            ["Escala regional", "150", "R$ 297", "R$ 44.550", "R$ 1.200"],
        ],
        "widths": [1800, 1700, 1900, 1900, 2060],
    },
}


MARKDOWN = """# Modelo de Negocio - Nova Forma CRM

**Versao:** 1.0
**Data:** Julho de 2026
**Objetivo:** estruturar o posicionamento, monetizacao, operacao e crescimento do Nova Forma CRM como produto SaaS vertical para construtoras.

## 1. Sumario executivo

O Nova Forma CRM deve ser posicionado como um CRM vertical para construtoras, empresas de steel frame, reformas e obras sob medida que recebem demanda pelo WhatsApp, site, Google Meu Negocio, Instagram e indicacoes.

A tese central e vender resultado comercial: menos leads perdidos, mais visitas marcadas, mais orcamentos enviados e maior controle da carteira de oportunidades.

## 2. Posicionamento

**Frase curta:** CRM comercial para construtoras que vendem pelo WhatsApp.

**Promessa:** pare de perder leads de obra e transforme conversas em visitas, orcamentos e contratos.

**Diferencial:** produto vertical, com funil, campos, IA, WhatsApp, tarefas e parceiro tecnico pensados para construcao.

## 3. Cliente ideal

- Construtoras pequenas e medias.
- Empresas de steel frame e construcao a seco.
- Empresas que fazem projeto + execucao.
- Operacoes que recebem leads pelo WhatsApp e ainda usam planilhas.
- Negocios que precisam de visita tecnica antes de orcar.

## 4. Problemas resolvidos

- Lead chega pelo WhatsApp e nao vira cadastro.
- Follow-up depende da memoria do vendedor.
- Dono nao enxerga pipeline real.
- Visita tecnica nao gera retorno estruturado.
- Cliente pede preco, visita ou orcamento, mas o proximo passo se perde.

## 5. Oferta comercial

Oferta recomendada: **implantacao assistida + assinatura mensal**.

Mensagem: **Em 7 dias organizamos seu atendimento comercial: funil, leads, follow-ups, templates de WhatsApp, dashboard e rotina de visitas.**

## 6. Planos sugeridos

| Plano | Preco sugerido | Publico | Inclui |
|---|---:|---|---|
| Essencial | R$ 97/mes | Autonomos e construtoras pequenas | 1 usuario, funil, tarefas, leads limitados, templates e dashboard basico |
| Profissional | R$ 197/mes | Construtoras em operacao comercial ativa | 3 usuarios, leads ilimitados, IA, interacoes, parceiro, briefing e dashboard completo |
| Empresa | R$ 397/mes | Equipes maiores e operacoes com parceiros | 10 usuarios, permissoes, parceiros, suporte prioritario e relatorios avancados |
| Setup | R$ 497 a R$ 1.500 | Todos os planos | Configuracao, importacao, treinamento, funil e templates iniciais |

## 7. Jornada de valor

| Etapa | Dor atual | Como o CRM resolve | Resultado esperado |
|---|---|---|---|
| Chegada do lead | Contato fica perdido no WhatsApp | IA ou cadastro manual cria lead estruturado | Lead entra no funil rapidamente |
| Qualificacao | Faltam cidade, terreno, planta e prazo | Campos obrigatorios e perguntas sugeridas | Atendimento mais consultivo |
| Follow-up | Retorno depende da memoria do vendedor | Tarefas e alertas de atraso | Menos oportunidades esquecidas |
| Visita tecnica | Briefing fica solto | Resumo do lead e retorno do parceiro | Orcamento com mais contexto |
| Orcamento | Cliente some apos proposta | Templates e follow-ups por etapa | Melhor taxa de conversao |

## 8. Estrategia comercial

1. Comecar com nicho de steel frame e construcao a seco.
2. Vender por demonstracao pratica usando conversas reais anonimizadas.
3. Cobrar setup para configurar o processo do cliente.
4. Medir sucesso por leads cadastrados, visitas marcadas e orcamentos enviados.
5. Expandir para reformas, casas modulares e construtoras tradicionais.

## 9. Operacao

Onboarding recomendado:

1. Diagnostico comercial.
2. Importacao dos leads existentes.
3. Configuracao de funil e status.
4. Cadastro de templates.
5. Treinamento do usuario admin.
6. Configuracao do parceiro tecnico.
7. Revisao depois de 15 dias.

## 10. Metricas

Metricas do cliente:

- Total de leads.
- Leads quentes.
- Leads atrasados.
- Visitas marcadas.
- Orcamentos enviados.
- Obras fechadas.
- Taxa de conversao.
- Valor potencial total.

Metricas do SaaS:

- MRR.
- Churn.
- CAC.
- LTV.
- Ativacao.
- Uso semanal.
- Receita de setup.

## 11. Riscos

- Baixa disciplina de uso.
- Concorrencia de CRMs genericos.
- Dependencia excessiva de IA.
- Excesso de customizacao.
- Suporte operacional caro.

## 12. Proximos passos

1. Criar uma pagina comercial simples.
2. Definir proposta de setup + mensalidade.
3. Vender para 5 clientes piloto.
4. Medir uso e conversao.
5. Ajustar planos.
6. Criar playbook de onboarding.
7. Expandir canais de venda.
"""


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_docx() -> None:
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.text = "Nova Forma CRM | Modelo de Negocio"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Documento estrategico - uso interno e comercial"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = GRAY

    add_paragraph(doc, "MODELO DE NEGOCIO", bold=True, color=GOLD, size=10, after=4)
    title = add_paragraph(doc, "Nova Forma CRM", bold=True, color=NAVY, size=28, after=6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph(doc, "Plano robusto para transformar o CRM em um SaaS vertical para construtoras, steel frame e atendimento comercial via WhatsApp.", color=GRAY, size=13, after=16)
    add_table(doc, ["Campo", "Definicao"], [
        ["Versao", "1.0"],
        ["Data", "Julho de 2026"],
        ["Uso", "Estrategia comercial, proposta de valor, operacao e crescimento"],
        ["Produto", "CRM web/mobile com Supabase, IA, funil, tarefas, WhatsApp, parceiros e briefing de visita"],
    ], widths=[1900, 7460], header_fill=PALE_GOLD)
    add_callout(doc, "Tese principal", "O produto nao deve ser vendido como CRM generico. Deve ser vendido como sistema comercial pronto para construtoras que querem parar de perder leads de obra no WhatsApp.", PALE_GOLD)

    doc.add_page_break()
    add_paragraph(doc, "Indice executivo", style="Heading 1")
    for item in [
        "Sumario executivo",
        "Produto e proposta de valor",
        "Cliente ideal e segmentacao",
        "Problemas resolvidos",
        "Diferenciais competitivos",
        "Oferta comercial",
        "Modelo de receita e precificacao",
        "Go-to-market",
        "Operacao e entrega",
        "Metricas de gestao",
        "Riscos e mitigacoes",
        "Roadmap estrategico",
        "Plano de 90 dias",
        "Anexos comerciais",
    ]:
        add_paragraph(doc, item, style=None, after=3)

    for title_text, paragraphs in SECTIONS:
        add_paragraph(doc, title_text, style="Heading 1")
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
        if title_text.startswith("6."):
            add_table(doc, TABLES["pricing"]["headers"], TABLES["pricing"]["rows"], TABLES["pricing"]["widths"], PALE_GOLD)
        if title_text.startswith("7."):
            add_table(doc, TABLES["financial"]["headers"], TABLES["financial"]["rows"], TABLES["financial"]["widths"], PALE_BLUE)
        if title_text.startswith("8."):
            add_table(doc, TABLES["journey"]["headers"], TABLES["journey"]["rows"], TABLES["journey"]["widths"], PALE_BLUE)

    add_paragraph(doc, "13. Plano de 90 dias", style="Heading 1")
    add_table(doc, ["Periodo", "Objetivo", "Entregaveis", "Indicador"], [
        ["Dias 1-15", "Preparar venda piloto", "Pagina comercial, demo, proposta e contrato simples", "5 reunioes marcadas"],
        ["Dias 16-30", "Fechar primeiros pilotos", "Setup assistido e importacao de leads", "3 clientes pagantes"],
        ["Dias 31-60", "Provar uso real", "Rotina semanal, ajustes de onboarding e templates", "70% dos leads com proxima acao"],
        ["Dias 61-90", "Repetir aquisicao", "Playbook de venda, indicacoes e cases", "10 clientes pagantes"],
    ], widths=[1400, 2300, 3700, 1960], header_fill=PALE_GOLD)

    add_paragraph(doc, "14. Anexos comerciais", style="Heading 1")
    add_paragraph(doc, "Pitch de 30 segundos", style="Heading 2")
    add_callout(doc, "Pitch", "O Nova Forma CRM organiza os leads que chegam pelo WhatsApp, transforma conversas em funil comercial, cria follow-ups, gera mensagens prontas e ajuda construtoras a marcar mais visitas e enviar mais orcamentos sem perder oportunidades.", LIGHT)
    add_paragraph(doc, "Checklist de qualificacao de cliente", style="Heading 2")
    for item in [
        "Recebe leads pelo WhatsApp toda semana.",
        "Tem ticket medio alto e venda consultiva.",
        "Precisa de visita tecnica antes do orcamento.",
        "Usa planilha ou memoria do vendedor para acompanhar clientes.",
        "Tem interesse em organizar atendimento e medir conversao.",
    ]:
        bullet = doc.add_paragraph(item, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(1)
    add_paragraph(doc, "Objeções e respostas", style="Heading 2")
    add_table(doc, ["Objecao", "Resposta sugerida"], [
        ["Ja uso WhatsApp", "O CRM nao substitui o WhatsApp; ele impede que os contatos se percam depois da conversa."],
        ["Nao tenho tempo para cadastrar", "A IA ajuda a transformar conversa ou print em rascunho de lead, e o usuario apenas revisa."],
        ["CRM e complicado", "O sistema ja vem com fluxo de obra, campos, status e mensagens prontas."],
        ["Quero testar antes", "Comece com um piloto de 30 dias com seus leads reais e acompanhe visitas e follow-ups."],
    ], widths=[2500, 6860], header_fill=PALE_BLUE)

    doc.save(DOCX_PATH)


def build_markdown() -> None:
    DOCS.mkdir(exist_ok=True)
    MD_PATH.write_text(MARKDOWN, encoding="utf-8")


if __name__ == "__main__":
    build_markdown()
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)
